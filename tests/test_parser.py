"""
Test script for document parser.

Creates sample documents and tests parsing functionality.
"""

from audiobook_utils import DocumentParser, estimate_processing_time
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import os


def create_sample_pdf(filename="test_sample.pdf"):
    """Create a sample PDF for testing."""
    c = canvas.Canvas(filename, pagesize=letter)
    
    # Page 1
    c.drawString(100, 750, "Chapter 1: The Beginning")
    c.drawString(100, 700, "This is a test PDF document for the audiobook converter.")
    c.drawString(100, 680, "It contains multiple pages and paragraphs to test the parsing functionality.")
    c.showPage()
    
    # Page 2
    c.drawString(100, 750, "Chapter 2: The Middle")
    c.drawString(100, 700, "This is the second page of our test document.")
    c.drawString(100, 680, "We want to ensure that multi-page documents are handled correctly.")
    c.showPage()
    
    c.save()
    print(f"✓ Created sample PDF: {filename}")


def create_sample_docx(filename="test_sample.docx"):
    """Create a sample DOCX for testing."""
    doc = Document()
    
    doc.add_heading('Chapter 1: The Beginning', level=1)
    doc.add_paragraph('This is a test DOCX document for the audiobook converter.')
    doc.add_paragraph('It contains multiple paragraphs to test the parsing functionality.')
    
    doc.add_heading('Chapter 2: The Middle', level=1)
    doc.add_paragraph('This is the second chapter of our test document.')
    doc.add_paragraph('We want to ensure that DOCX documents are handled correctly.')
    
    doc.save(filename)
    print(f"✓ Created sample DOCX: {filename}")


def test_parser():
    """Test the document parser with sample files."""
    parser = DocumentParser()
    
    print("\n" + "="*60)
    print("DOCUMENT PARSER TEST")
    print("="*60)
    
    # Test PDF
    print("\n📄 Testing PDF parsing...")
    try:
        text, metadata = parser.parse_document("test_sample.pdf")
        print(f"✓ PDF parsed successfully!")
        print(f"  Pages: {metadata['page_count']}")
        print(f"  Words: {metadata['word_count']}")
        print(f"  Characters: {metadata['char_count']}")
        print(f"\n  First 200 chars:\n  {text[:200]}...")
        
        # Test chunking
        chunks = parser.chunk_text(text, max_chars=100)
        print(f"\n✓ Text chunked into {len(chunks)} chunks")
        print(f"  First chunk: {chunks[0][:80]}...")
        
        # Estimate processing time
        mins, secs = estimate_processing_time(metadata['char_count'])
        print(f"\n⏱️  Estimated processing time: {mins}m {secs}s")
        
    except Exception as e:
        print(f"✗ PDF test failed: {e}")
    
    # Test DOCX
    print("\n📝 Testing DOCX parsing...")
    try:
        text, metadata = parser.parse_document("test_sample.docx")
        print(f"✓ DOCX parsed successfully!")
        print(f"  Estimated pages: {metadata['page_count']}")
        print(f"  Words: {metadata['word_count']}")
        print(f"  Characters: {metadata['char_count']}")
        print(f"\n  First 200 chars:\n  {text[:200]}...")
        
        # Test chunking
        chunks = parser.chunk_text(text, max_chars=100)
        print(f"\n✓ Text chunked into {len(chunks)} chunks")
        
        # Estimate processing time
        mins, secs = estimate_processing_time(metadata['char_count'])
        print(f"\n⏱️  Estimated processing time: {mins}m {secs}s")
        
    except Exception as e:
        print(f"✗ DOCX test failed: {e}")
    
    print("\n" + "="*60)
    print("✅ ALL TESTS COMPLETED")
    print("="*60 + "\n")


if __name__ == "__main__":
    # Create sample documents
    create_sample_pdf()
    create_sample_docx()
    
    # Test parser
    test_parser()
    
    # Cleanup
    print("🧹 Cleaning up test files...")
    if os.path.exists("test_sample.pdf"):
        os.remove("test_sample.pdf")
    if os.path.exists("test_sample.docx"):
        os.remove("test_sample.docx")
    print("✓ Cleanup complete")
