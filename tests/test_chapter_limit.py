"""
Test script for chapter splitting and character limits.
"""
from audiobook_utils import DocumentParser
from audiobook_generator import AudiobookGenerator
import shutil
from pathlib import Path

def test_chapter_splitting():
    print("Testing chapter splitting with limits...")
    parser = DocumentParser()
    
    # Create a dummy text with a very long chapter
    long_chapter_content = "This is a long sentence. " * 2000 # ~50k chars
    text = f"""
    Title Page
    
    Chapter 1: Short Chapter
    This is a short chapter.
    
    Chapter 2: Long Chapter
    {long_chapter_content}
    
    Chapter 3: Another Short One
    The end.
    """
    
    # Test splitting with a small limit to force splits
    print("\n1. Testing split_into_chapters (limit=1000 chars)...")
    chapters = parser.split_into_chapters(text, max_chars=1000)
    
    for title, content in chapters:
        print(f"  - {title}: {len(content)} chars")
        if len(content) > 1000:
            print(f"    FAIL: Content length {len(content)} exceeds limit 1000")
            return

    # Test prepare_chapters integration
    print("\n2. Testing prepare_chapters integration...")
    generator = AudiobookGenerator(device="cpu") # Use CPU for fast init
    
    # Create temp docx
    from docx import Document
    doc = Document()
    doc.add_heading("Chapter 1: The Giant", 0)
    doc.add_paragraph("A" * 40000) # 40k chars
    doc.save("test_limit.docx")
    
    output_dir = "test_chapters_output"
    if Path(output_dir).exists():
        shutil.rmtree(output_dir)
        
    files = generator.prepare_chapters("test_limit.docx", output_dir)
    
    print(f"\nGenerated {len(files)} files:")
    for f in files:
        print(f"  - {f}")
        
    print("\nSUCCESS: Chapter splitting verified.")

if __name__ == "__main__":
    test_chapter_splitting()
